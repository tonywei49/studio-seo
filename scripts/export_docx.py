#!/usr/bin/env python3
from __future__ import annotations

import json
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Pt


def set_run_font(run, size_pt: int, bold: bool = False) -> None:
    run.bold = bold
    run.font.size = Pt(size_pt)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def add_paragraph(document: Document, text: str, size_pt: int, bold: bool = False, space_after_pt: int = 0) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(space_after_pt)
    run = paragraph.add_run(text)
    set_run_font(run, size_pt, bold)


def add_body_block(document: Document, heading: str, body: str) -> None:
    if not body.strip():
        return
    add_paragraph(document, heading, 14, True, 8)
    paragraphs = [item.strip() for item in body.replace("\r\n", "\n").split("\n") if item.strip()]
    for index, item in enumerate(paragraphs):
        add_paragraph(document, item, 14, False, 10)
        if index != len(paragraphs) - 1:
            document.add_paragraph()


def add_tdk_block(document: Document, label: str, value: str) -> None:
    if not value.strip():
        return
    add_paragraph(document, f"{label}{value}", 14, False, 10)


def main() -> int:
    if len(sys.argv) != 3:
        raise SystemExit("usage: export_docx.py <payload.json> <output.docx>")

    payload_path = Path(sys.argv[1])
    output_path = Path(sys.argv[2])
    payload = json.loads(payload_path.read_text(encoding="utf-8"))

    document = Document()
    normal_style = document.styles["Normal"]
    normal_style.font.name = "Microsoft YaHei"
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal_style.font.size = Pt(14)

    output_language = payload["outputLanguage"]
    title_zh = payload["selectedTitleZh"]
    title_en = payload["selectedTitleEn"]

    add_paragraph(document, title_en if output_language == "en" else title_zh, 20, True, 12)
    if output_language == "zh-en":
        add_paragraph(document, title_en, 20, True, 14)

    if output_language != "en":
        add_body_block(document, "中文正文", payload["bodyZh"])
        document.add_paragraph()

    if output_language != "zh":
        add_body_block(document, "English Body", payload["bodyEn"])
        document.add_paragraph()

    add_paragraph(document, "TDK", 14, True, 8)
    if output_language != "en":
        add_tdk_block(document, "Title (ZH): ", payload["tdkTitleZh"])
        add_tdk_block(document, "Description (ZH): ", payload["tdkDescriptionZh"])
        add_tdk_block(document, "Keywords (ZH): ", payload["tdkKeywordsZh"])
        document.add_paragraph()
    if output_language != "zh":
        add_tdk_block(document, "Title (EN): ", payload["tdkTitleEn"])
        add_tdk_block(document, "Description (EN): ", payload["tdkDescriptionEn"])
        add_tdk_block(document, "Keywords (EN): ", payload["tdkKeywordsEn"])

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
